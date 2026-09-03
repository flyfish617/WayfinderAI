from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path(r"E:\develop\agent开发\WayfinderAI")
SOURCE = ROOT / "刘玉钦-27届-后端实习简历-智能旅行助手.docx"
OUTPUT = ROOT / "刘玉钦-27届-后端实习简历-智能旅行助手-含实习经历.docx"


def clear_runs(paragraph):
    for run in list(paragraph.runs):
        paragraph._p.remove(run._r)


def configure_run(run, bold=False):
    run.bold = bold
    run.font.name = "微软雅黑"
    run.font.size = Pt(10)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    for key in ("ascii", "hAnsi", "eastAsia"):
        rfonts.set(qn(f"w:{key}"), "微软雅黑")


def set_segments(paragraph, segments):
    clear_runs(paragraph)
    for text, bold in segments:
        run = paragraph.add_run(text)
        configure_run(run, bold=bold)


doc = Document(SOURCE)
paragraphs = doc.paragraphs
if len(paragraphs) < 15:
    raise RuntimeError("简历段落数量异常，停止修改")
if paragraphs[10].text.strip() != "重庆民航凯亚信息技术有限公司\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t2026.6-至今":
    # Allow Word to normalize tabs while still checking the company/date anchor.
    if "重庆民航凯亚信息技术有限公司" not in paragraphs[10].text or "2026.6-至今" not in paragraphs[10].text:
        raise RuntimeError("实习经历锚点不一致，停止修改")
if paragraphs[11].text.strip() != "AI全栈开发实习  |  航际通综合服务平台":
    if "航际通综合服务平台" not in paragraphs[11].text:
        raise RuntimeError("实习项目锚点不一致，停止修改")

set_segments(paragraphs[12], [
    ("• 监控模块设计与拆解：", True),
    ("参与监控模块协作设计，输出功能点汇总与设计文档，拆解运行看板、任务协同、项目管理、异常预警，明确指标口径、角色权限和预警闭环。", False),
])

set_segments(paragraphs[13], [
    ("• 异常预警全链路：", True),
    ("独立完成异常预警前后端开发；实现规则、接收人、阈值、级别与静默期管理，接入 Feign 工单指标、OSHI 服务器指标及 HertzBeat 规则同步/回调。", False),
])

set_segments(paragraphs[14], [
    ("• 通知与处置闭环：", True),
    ("实现 WebSocket 定向弹窗、异步邮件、预警记录与统计；以已读→处理中→已解决→已结束状态机和权限校验保障预警处理闭环。", False),
])

# 为加入实习经历后的单页版面，将旅行助手亮点合并为两条高密度要点。
set_segments(paragraphs[26], [
    ("• Agent 流编排与分层记忆：", True),
    ("设计 Memory RAG→MCP→Planner Agent 链路，FAISS 召回用户偏好/旅行知识，ContextManager 聚合上下文，Pydantic 校验 LLM JSON。", False),
])
set_segments(paragraphs[27], [
    ("• MCP 工具与异步治理：", True),
    ("高德 MCP 单会话+Redis TTL 缓存；Redis 队列 Worker 承载长任务，版本快照/乐观锁支持回滚。", False),
])
for paragraph in (paragraphs[28], paragraphs[29]):
    paragraph._element.getparent().remove(paragraph._element)

doc.save(OUTPUT)
print(OUTPUT)
